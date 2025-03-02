"""
Utility functions for extracting text from various document formats.
"""
import os
from typing import Dict, List, Optional, Tuple

import docx
import pdfplumber
import PyPDF2


def extract_text_from_pdf(file_path: str) -> Tuple[str, Dict]:
    """
    Extract text from a PDF file using PyPDF2 and pdfplumber.

    Args:
        file_path: Path to the PDF file

    Returns:
        Tuple containing:
        - Extracted text as a string
        - Metadata dictionary with page count and other information
    """
    text = ""
    metadata = {"page_count": 0, "pages": {}}

    # Try with PyPDF2 first
    try:
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            metadata["page_count"] = len(reader.pages)

            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                text += page_text + "\n\n"
                metadata["pages"][i] = {"char_count": len(page_text)}

    except Exception as e:
        print(f"PyPDF2 extraction failed: {e}")

        # Fallback to pdfplumber if PyPDF2 fails
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                metadata["page_count"] = len(pdf.pages)

                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    text += page_text + "\n\n"
                    metadata["pages"][i] = {"char_count": len(page_text)}

        except Exception as e2:
            print(f"pdfplumber extraction also failed: {e2}")
            return f"Error extracting text: {e}, {e2}", {"error": str(e)}

    return text, metadata


def extract_text_from_docx(file_path: str) -> Tuple[str, Dict]:
    """
    Extract text from a DOCX file.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Tuple containing:
        - Extracted text as a string
        - Metadata dictionary with paragraph count and other information
    """
    try:
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        metadata = {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "section_count": len(doc.sections),
        }

        return text, metadata

    except Exception as e:
        print(f"DOCX extraction failed: {e}")
        return f"Error extracting text: {e}", {"error": str(e)}


def extract_text_from_file(file_path: str) -> Tuple[str, Dict]:
    """
    Extract text from a file based on its extension.

    Args:
        file_path: Path to the file

    Returns:
        Tuple containing:
        - Extracted text as a string
        - Metadata dictionary with file-specific information
    """
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()

    if file_extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif file_extension in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    else:
        return f"Unsupported file format: {file_extension}", {
            "error": f"Unsupported format: {file_extension}"
        }


def get_file_metadata(file_path: str) -> Dict:
    """
    Get basic metadata about a file.

    Args:
        file_path: Path to the file

    Returns:
        Dictionary containing file metadata
    """
    try:
        file_stats = os.stat(file_path)
        file_name = os.path.basename(file_path)
        _, file_extension = os.path.splitext(file_name)

        metadata = {
            "file_name": file_name,
            "file_extension": file_extension.lower(),
            "file_size_bytes": file_stats.st_size,
            "file_size_kb": round(file_stats.st_size / 1024, 2),
            "last_modified": file_stats.st_mtime,
        }

        return metadata

    except Exception as e:
        print(f"Error getting file metadata: {e}")
        return {"error": str(e)}
